#!/usr/bin/env python3
"""
PRD 초안 .docx 생성기 (AI2W 퍼실리테이터)

사용법:
    python3 build_prd_docx.py <prd_data.json> <output.docx>

입력 JSON 형식:
{
  "team": "○○팀",
  "date": "2026.00.00",
  "agent_name": "○○ Agent",
  "value": "한 줄 가치 제안",
  "sections": {
    "1_overview": "...",            // 개요(이름/목적/가치)
    "2_context": "...",            // 업무 맥락(As-Is + Pain Point 3)
    "3_users": "...",              // ★대상 사용자
    "4_why": "...",                // 선정 이유
    "5_requirements": "...",       // ★기능 요구사항
    "6_human": "...",              // Human 개입 설계
    "7_implementation": "...",     // ★구현 방향
    "8_impact": "...",             // 기대효과 & 측정지표
    "9_risk": "..."                // 리스크 & 전제
  }
}
값은 문자열 또는 문자열 리스트(각 항목이 bullet)로 줄 수 있다.

python-docx 필요:  pip install python-docx
"""
import json
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx가 필요합니다.  실행:  pip install python-docx")

# 디자인 색 팔레트 (Jacksons Purple + Gray)
UPS = RGBColor(0x80, 0x5C, 0xFB)      # Ups Purple — 포인트 액센트 전용
JP400 = RGBColor(0x75, 0x78, 0xFF)    # 포인트 텍스트
GRAY700 = RGBColor(0x0D, 0x13, 0x20)  # 최고 어두운 제목
GRAY600 = RGBColor(0x1C, 0x25, 0x37)  # 제목
GRAY500 = RGBColor(0x43, 0x4C, 0x60)  # 본문
GRAY400 = RGBColor(0x97, 0x9C, 0xAE)  # 캡션/각주

SECTIONS = [
    ("1_overview", "1. 개요", False),
    ("2_context", "2. 업무 맥락 (As-Is + Pain Point)", False),
    ("3_users", "3. 대상 사용자", True),
    ("4_why", "4. 선정 이유", False),
    ("5_requirements", "5. 기능 요구사항", True),
    ("6_human", "6. Human 개입 설계", False),
    ("7_implementation", "7. 구현 방향", True),
    ("8_impact", "8. 기대효과 & 측정지표", False),
    ("9_risk", "9. 리스크 & 전제", False),
]


def set_font(run, name="Noto Sans KR", size=11, color=None, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 한글 폰트 적용
    try:
        from docx.oxml.ns import qn
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass


def add_body(doc, value):
    """문자열 또는 리스트를 본문으로 추가."""
    items = value if isinstance(value, list) else [value]
    for it in items:
        it = (it or "").strip()
        if not it:
            it = "(추후 확인)"
        p = doc.add_paragraph()
        if isinstance(value, list):
            p.style = doc.styles["List Bullet"]
        r = p.add_run(it)
        set_font(r, size=11, color=GRAY500)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5


def main():
    if len(sys.argv) < 3:
        sys.exit("사용법: python3 build_prd_docx.py <prd_data.json> <output.docx>")
    data = json.load(open(sys.argv[1], encoding="utf-8"))
    out = sys.argv[2]
    sec = data.get("sections", {})

    doc = Document()
    doc.sections[0].left_margin = Inches(1.0)
    doc.sections[0].right_margin = Inches(1.0)

    # 제목
    label = doc.add_paragraph()
    r = label.add_run("AI2W WORKSHOP · PRD")
    set_font(r, name="Geist", size=9, color=UPS, bold=True)
    label.paragraph_format.space_after = Pt(2)
    h = doc.add_paragraph()
    r = h.add_run("PRD 초안")
    set_font(r, size=24, color=GRAY700, bold=True)
    h.paragraph_format.space_after = Pt(2)
    sub = doc.add_paragraph()
    r = sub.add_run(data.get("agent_name", "AI Agent"))
    set_font(r, size=15, color=UPS, bold=True)   # 시그니처 퍼플 액센트
    if data.get("value"):
        v = doc.add_paragraph()
        r = v.add_run(data["value"])
        set_font(r, size=11, color=GRAY400)
    meta = doc.add_paragraph()
    r = meta.add_run(f"팀: {data.get('team','')}    작성일: {data.get('date','')}    "
                     f"AI2W Workshop")
    set_font(r, size=9, color=GRAY400)
    doc.add_paragraph()

    # 섹션
    for key, title, core in SECTIONS:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(10)
        r = hp.add_run(title)
        set_font(r, size=13, color=GRAY600, bold=True)
        if core:
            r2 = hp.add_run("   ★ 핵심")
            set_font(r2, size=11, color=UPS, bold=True)
        add_body(doc, sec.get(key, ""))

    # 푸터 안내
    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run("★ 핵심 3칸(대상 사용자·기능 요구사항·구현 방향)은 실사용을 좌우하는 항목입니다.")
    set_font(r, size=9, color=GRAY400)

    doc.save(out)
    print("WROTE", out)


if __name__ == "__main__":
    main()
